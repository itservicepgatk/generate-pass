"""Сборка Word-документа — точные размеры + отступы для резки"""

import io
from PIL import Image
from docx import Document
from docx.shared import Cm, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ROW_HEIGHT_RULE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from config import PassConfig
from card_renderer import CardRenderer
from photo_utils import PhotoUtils


class DocumentBuilder:

    CHUNK = 8

    def __init__(self, cfg: PassConfig):
        self.cfg = cfg
        self.renderer = CardRenderer(cfg)

    def build(
        self,
        photos: dict[str, bytes],
        logo_bytes: bytes | None = None,
        progress_cb=None,
    ) -> bytes:
        logo_pil = None
        if logo_bytes:
            logo_pil = Image.open(io.BytesIO(logo_bytes)).convert("RGBA")

        doc = self._new_doc()
        names = list(photos.keys())
        chunks = [names[i:i + self.CHUNK] for i in range(0, len(names), self.CHUNK)]
        total = len(names)
        done = 0

        for ci, chunk in enumerate(chunks):
            if ci > 0:
                doc.add_page_break()
            tf = self._table(doc)

            for i, fio in enumerate(chunk):
                photo_pil = PhotoUtils.process_upload(photos[fio], fio)
                card = self.renderer.front(photo_pil, logo_pil)
                self._insert(tf, i // 2, i % 2, card)
                done += 1
                if progress_cb:
                    progress_cb(done / (total * 2))

            doc.add_page_break()
            tb = self._table(doc)

            for i, fio in enumerate(chunk):
                card = self.renderer.back(fio)
                self._insert(tb, i // 2, 1 - (i % 2), card)
                done += 1
                if progress_cb:
                    progress_cb(done / (total * 2))

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()

    # ── Приватные ──────────────────────────────────

    def _new_doc(self):
        doc = Document()
        s = doc.sections[0]
        for attr in ("left_margin", "right_margin", "top_margin", "bottom_margin"):
            setattr(s, attr, Mm(10))
        s.page_width, s.page_height = Mm(210), Mm(297)
        return doc

    def _table(self, doc):
        t = doc.add_table(rows=4, cols=2)
        t.autofit = False
        t.allow_autofit = False
        t.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # ══ Ширина колонки = карточка + отступ для резки ══
        margin = self.cfg.cut_margin
        col_w = Cm(self.cfg.card_w + margin * 2)
        for c in t.columns:
            c.width = col_w

        # ══ Небольшие отступы ячеек для зазора между карточками ══
        self._set_cell_margins(t, margin)
        self._remove_table_borders(t)

        return t

    def _insert(self, table, row, col, card_img: Image.Image):
        buf = io.BytesIO()
        card_img.save(buf, format="PNG")
        buf.seek(0)

        margin = self.cfg.cut_margin
        cell = table.rows[row].cells[col]
        cell.width = Cm(self.cfg.card_w + margin * 2)

        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Убираем лишние отступы параграфа
        pf = p.paragraph_format
        pf.space_before = Cm(0)
        pf.space_after = Cm(0)

        # ══ Картинка = ТОЧНО размер карточки ══
        p.add_run().add_picture(
            buf,
            width=Cm(self.cfg.card_w),
            height=Cm(self.cfg.card_h),
        )

        # ══ Высота строки = карточка + отступы сверху/снизу ══
        table.rows[row].height_rule = WD_ROW_HEIGHT_RULE.EXACTLY
        table.rows[row].height = Cm(self.cfg.card_h + margin * 2)

    # ── XML-утилиты ────────────────────────────────

    @staticmethod
    def _set_cell_margins(table, margin_cm: float):
        """
        Устанавливает отступы ячеек таблицы.
        margin_cm — отступ в сантиметрах с каждой стороны.
        """
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # Удаляем старые
        for old in tblPr.findall(qn('w:tblCellMar')):
            tblPr.remove(old)

        # 1 см = 567 twips (единица измерения Word)
        twips = int(margin_cm * 567)

        tcMar = OxmlElement('w:tblCellMar')
        for side in ('top', 'left', 'bottom', 'right'):
            node = OxmlElement(f'w:{side}')
            node.set(qn('w:w'), str(twips))
            node.set(qn('w:type'), 'dxa')
            tcMar.append(node)
        tblPr.append(tcMar)

    @staticmethod
    def _remove_table_borders(table):
        """Убирает видимые рамки таблицы (карточки и так с рамками)"""
        tbl = table._tbl
        tblPr = tbl.tblPr
        if tblPr is None:
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        for old in tblPr.findall(qn('w:tblBorders')):
            tblPr.remove(old)

        borders = OxmlElement('w:tblBorders')
        for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
            b = OxmlElement(f'w:{side}')
            b.set(qn('w:val'), 'none')
            b.set(qn('w:sz'), '0')
            b.set(qn('w:space'), '0')
            b.set(qn('w:color'), 'auto')
            borders.append(b)
        tblPr.append(borders)